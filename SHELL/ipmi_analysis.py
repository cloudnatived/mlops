from docx.shared import Cm, Pt, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx import Document
from docx.oxml.ns import qn
from docx.shared import Inches
from datetime import datetime
import os
import sys
import re
from collections import Counter
import matplotlib.pyplot as plt
import tempfile

# 设置 matplotlib 支持中文显示
plt.rcParams['font.sans-serif'] = ['SimHei']  # 使用黑体
plt.rcParams['axes.unicode_minus'] = False  # 解决负号显示问题

# 全局文档对象
document = Document()
patchs = os.path.dirname(os.path.abspath(__file__))
t1 = datetime.now().strftime('%Y-%m-%d')

def yemei_tupian():
    """添加页眉图片"""
    global document
    header = document.sections[0].header
    paragraph = header.paragraphs[0]  # Use first paragraph
    paragraph.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = paragraph.add_run()
    try:
        run.add_picture(os.path.join(patchs, 'img/yemei.png'), width=Cm(15))
        run.underline = True
    except FileNotFoundError:
        print("警告: 页眉图片未找到，将使用默认文本页眉")
        run.add_text("IPMI巡检报告")

def top_diyiye(report_title):
    """生成报告首页"""
    global document
    topbt = f"{report_title} IPMI日志巡检评估报告"
    
    heading = document.add_heading(topbt, level=1)
    heading.style.font.size = Pt(24)
    heading.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    for _ in range(7):
        document.add_paragraph('\n')
    
    document.add_heading('北京银信长远科技股份有限公司', 2).paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    document.add_paragraph(t1).paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    document.add_page_break()
    document.add_heading('巡检目的', 2)
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.first_line_indent = Pt(25)
    run = paragraph.add_run(
        '为保障数据中心物理服务器，业务系统稳定高效的运行，针对IPMI日志、设备运行信息进行巡检，'
        '发现设备运行中可能出现及已经存在的问题，减少和避免设备因为配置问题及硬件存在的故障产生事故，'
        '给客户IT信息系统运维工作提供有力的参考依据。'
    )
    run.font.name = '宋体'
    run.font.size = Pt(12)
    r = run._element
    r.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

def parse_ipmi_log(log_path):
    """解析IPMI日志文件，提取关键信息，从多个维度分析"""
    log_data = {
        'device_ip': '未知',
        'system_info': {},
        'hardware_status': {'memory_slots': [], 'drive_slots': [], 'processor_status': []},
        'sel_info': {},
        'sel_events': {'total': 0, 'critical_events': [], 'power_state_changes': [], 'memory_errors': [], 'processor_errors': []},
        'fru_info': {},
        'network_info': {},
        'security_info': {},
        'system_resources': {'memory': {}, 'disk': {}, 'processes': []},
        'event_summary': {'by_type': {}, 'by_date': {}}
    }

    try:
        with open(log_path, 'r', encoding='utf-8') as f:
            log_content = f.read()
    except Exception as e:
        print(f"错误: 无法读取日志文件 {log_path}: {str(e)}")
        return log_data

    # Extract device IP from filename or log
    ip_match = re.search(r'(\d{1,3}\.\d{1,3}\.\d{1,3}\.\d{1,3})', log_path)
    if ip_match:
        log_data['device_ip'] = ip_match.group(1)

    sections = re.split(r'#+', log_content)
    section = None

    for section_content in sections:
        section_content = section_content.strip()
        if not section_content:
            continue

        # Identify section
        if 'mc info' in section_content:
            section = 'system_info'
        elif 'sel info' in section_content:
            section = 'sel_info'
        elif 'fru list' in section_content:
            section = 'fru_info'
        elif 'power status' in section_content:
            section = 'power'
        elif 'lan print' in section_content or 'ipmcget -d macaddr' in section_content:
            section = 'network'
        elif 'sel list' in section_content:
            section = 'sel_events'
        elif 'user list' in section_content:
            section = 'user_list'
        elif 'session info active' in section_content:
            section = 'session'
        elif 'free' in section_content:
            section = 'memory'
        elif 'df' in section_content:
            section = 'disk'
        elif 'ps' in section_content:
            section = 'processes'

        # Parse section data
        if section == 'system_info':
            for line in section_content.split('\n'):
                if ':' in line:
                    key, value = map(str.strip, line.split(':', 1))
                    log_data['system_info'][key] = value
        elif section == 'sel_info':
            for line in section_content.split('\n'):
                if ':' in line:
                    key, value = map(str.strip, line.split(':', 1))
                    log_data['sel_info'][key] = value
        elif section == 'fru_info':
            for line in section_content.split('\n'):
                if ':' in line:
                    key, value = map(str.strip, line.split(':', 1))
                    if key in ['Product Manufacturer', 'Product Name', 'Product Part Number', 'Product Serial', 'Board Manufacturer', 'Board Product Name', 'Board Serial Number']:
                        log_data['fru_info'][key] = value
        elif section == 'power':
            match = re.search(r'Chassis Power is (\w+)', section_content)
            if match:
                log_data['power_status'] = match.group(1)
        elif section == 'network':
            for line in section_content.split('\n'):
                if ':' in line:
                    key, value = map(str.strip, line.split(':', 1))
                    log_data['network_info'][key] = value
                if 'NIC1' in line:
                    mac_match = re.search(r'NIC1:([0-9A-Fa-f:-]+)', line)
                    if mac_match:
                        log_data['network_info']['MAC Address'] = mac_match.group(1)
        elif section == 'sel_events':
            for line in section_content.split('\n'):
                if re.match(r'\s*\w+\s*\|.*\|\s*(Asserted|Deasserted)', line):
                    log_data['sel_events']['total'] += 1
                    event_type = re.search(r'\| ([^|]+) \|', line).group(1).strip()
                    event_date = re.search(r'(\d{2}/\d{2}/\d{4}|\w+-\w+)', line)
                    event_date = event_date.group(1) if event_date else 'Unknown'
                    
                    # Update event summary
                    log_data['event_summary']['by_type'][event_type] = log_data['event_summary']['by_type'].get(event_type, 0) + 1
                    log_data['event_summary']['by_date'][event_date] = log_data['event_summary']['by_date'].get(event_date, 0) + 1

                    if 'Power Supply' in line and ('Failure detected' in line or 'Power Supply AC lost' in line):
                        log_data['sel_events']['critical_events'].append(line.strip())
                    if 'System ACPI Power State' in line:
                        log_data['sel_events']['power_state_changes'].append(line.strip())
                    if 'Memory #' in line and 'Configuration Error' in line:
                        log_data['sel_events']['memory_errors'].append(line.strip())
                    if 'Processor #' in line and 'IERR' in line:
                        log_data['sel_events']['processor_errors'].append(line.strip())
                    if 'Memory #' in line:
                        memory_match = re.search(r'Memory #0x(\w+)\s*\|\s*Presence Detected\s*\|\s*(Asserted|Deasserted)', line)
                        if memory_match:
                            slot, status = memory_match.groups()
                            log_data['hardware_status']['memory_slots'].append({'slot': f'Memory 0x{slot}', 'status': status})
                    if 'Drive Slot / Bay #' in line:
                        drive_match = re.search(r'Drive Slot / Bay #0x(\w+)\s*\|\s*Drive Present\s*\|\s*(Asserted|Deasserted)', line)
                        if drive_match:
                            slot, status = drive_match.groups()
                            log_data['hardware_status']['drive_slots'].append({'slot': f'Drive 0x{slot}', 'status': status})
                    if 'Processor #' in line:
                        proc_match = re.search(r'Processor #0x(\w+)\s*\|\s*Presence detected\s*\|\s*(Asserted|Deasserted)', line)
                        if proc_match:
                            slot, status = proc_match.groups()
                            log_data['hardware_status']['processor_status'].append({'slot': f'Processor 0x{slot}', 'status': status})
        elif section == 'user_list':
            users = []
            for line in section_content.split('\n'):
                user_match = re.search(r'(\d+)\s+([^\s]*)\s+true\s+true\s+true\s+([^\s]+)', line)
                if user_match:
                    user_id, name, priv = user_match.groups()
                    users.append({'id': user_id, 'name': name, 'privilege': priv})
            log_data['security_info']['users'] = users
        elif section == 'session':
            for line in section_content.split('\n'):
                if ':' in line:
                    key, value = map(str.strip, line.split(':', 1))
                    log_data['security_info'][key] = value
        elif section == 'memory':
            for line in section_content.split('\n'):
                if line.startswith('Mem:'):
                    parts = line.split()
                    log_data['system_resources']['memory'] = {
                        'total': parts[1],
                        'used': parts[2],
                        'free': parts[3]
                    }
        elif section == 'disk':
            for line in section_content.split('\n'):
                if line.startswith('/dev/') or line.startswith('tmpfs'):
                    parts = line.split()
                    log_data['system_resources']['disk'][parts[0]] = {
                        'total': parts[1],
                        'used': parts[2],
                        'available': parts[3],
                        'use%': parts[4]
                    }
        elif section == 'processes':
            for line in section_content.split('\n'):
                if line.startswith('  PID'):
                    continue
                proc_match = re.search(r'(\d+)\s+(\S+)\s+(\d+)\s+(\S+)\s+(.+)', line)
                if proc_match:
                    pid, user, vsz, stat, command = proc_match.groups()
                    log_data['system_resources']['processes'].append({
                        'pid': pid,
                        'user': user,
                        'vsz': vsz,
                        'stat': stat,
                        'command': command
                    })

    return log_data

def translate_event_type(event_type):
    """将事件类型翻译为中文"""
    translations = {
        'Add-in Card': '附加卡',
        'Memory': '内存',
        'Drive Slot / Bay': '驱动器插槽',
        'Processor': '处理器',
        'System Boot Initiated': '系统启动',
        'System ACPI Power State': '系统ACPI电源状态',
        'Power Supply': '电源供应'
    }
    return translations.get(event_type.split('#')[0].strip(), event_type)

def generate_report(log_data, inspector):
    """生成IPMI巡检报告（docx格式）"""
    global document
    device_ip = log_data['device_ip']
    document.add_heading(f'设备IP: {device_ip}', level=1)

    # 设置页面边距
    sections = document.sections
    for section in sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(3)
        section.right_margin = Cm(3)

    # 1. 系统信息表格
    document.add_heading('1. 系统基本信息', level=2)
    table = document.add_table(rows=8, cols=2)
    table.style = 'Table Grid'
    table.columns[0].width = Cm(5)
    table.columns[1].width = Cm(10)
    table.cell(0,0).text = '设备型号'
    table.cell(0,1).text = log_data['fru_info'].get('Product Name', '未知')
    table.cell(1,0).text = '固件版本'
    table.cell(1,1).text = log_data['system_info'].get('Firmware Revision', '未知')
    table.cell(2,0).text = 'IPMI版本'
    table.cell(2,1).text = log_data['system_info'].get('IPMI Version', '未知')
    table.cell(3,0).text = '电源状态'
    table.cell(3,1).text = log_data.get('power_status', '未知')
    table.cell(4,0).text = '产品制造商'
    table.cell(4,1).text = log_data['fru_info'].get('Product Manufacturer', '未知')
    table.cell(5,0).text = '产品部件编号'
    table.cell(5,1).text = log_data['fru_info'].get('Product Part Number', '未知')
    table.cell(6,0).text = '产品序列号'
    table.cell(6,1).text = log_data['fru_info'].get('Product Serial', '未知')
    table.cell(7,0).text = '主板制造商'
    table.cell(7,1).text = log_data['fru_info'].get('Board Manufacturer', '未知')

    # 2. 硬件信息表格
    document.add_heading('2. 硬件信息', level=2)
    table = document.add_table(rows=1, cols=3)
    table.style = 'Table Grid'
    table.columns[0].width = Cm(5)
    table.columns[1].width = Cm(5)
    table.columns[2].width = Cm(5)
    table.cell(0,0).text = '组件'
    table.cell(0,1).text = '插槽'
    table.cell(0,2).text = '状态'
    for component in ['memory_slots', 'drive_slots', 'processor_status']:
        for item in log_data['hardware_status'].get(component, []):
            row_cells = table.add_row().cells
            row_cells[0].text = component.replace('_slots', '').replace('_status', '').capitalize()
            row_cells[1].text = item['slot']
            row_cells[2].text = item['status']

    # 3. 网络信息
    document.add_heading('3. 网络配置', level=2)
    table = document.add_table(rows=3, cols=2)
    table.style = 'Table Grid'
    table.columns[0].width = Cm(5)
    table.columns[1].width = Cm(10)
    table.cell(0,0).text = 'IP地址'
    table.cell(0,1).text = log_data['network_info'].get('IP Address', log_data['device_ip'])
    table.cell(1,0).text = '子网掩码'
    table.cell(1,1).text = log_data['network_info'].get('Subnet Mask', '未知')
    table.cell(2,0).text = 'MAC地址'
    table.cell(2,1).text = log_data['network_info'].get('MAC Address', '未知')

    # 4. 系统资源使用情况
    document.add_heading('4. 系统资源使用情况', level=2)
    
    # 4.1 内存使用
    document.add_heading('4.1 内存使用情况', level=3)
    mem_info = log_data['system_resources'].get('memory', {})
    table = document.add_table(rows=4, cols=2)
    table.style = 'Table Grid'
    table.columns[0].width = Cm(5)
    table.columns[1].width = Cm(10)
    table.cell(0,0).text = '总内存'
    table.cell(0,1).text = mem_info.get('total', '未知') + ' KB'
    table.cell(1,0).text = '已使用'
    table.cell(1,1).text = mem_info.get('used', '未知') + ' KB'
    table.cell(2,0).text = '可用'
    table.cell(2,1).text = mem_info.get('free', '未知') + ' KB'
    table.cell(3,0).text = '使用率'
    if mem_info.get('total') and mem_info.get('used'):
        usage = (int(mem_info['used']) / int(mem_info['total'])) * 100
        table.cell(3,1).text = f'{usage:.1f}%'
    else:
        table.cell(3,1).text = '未知'

    # 4.2 磁盘使用
    document.add_heading('4.2 磁盘使用情况', level=3)
    disk_info = log_data['system_resources'].get('disk', {})
    table = document.add_table(rows=1, cols=5)
    table.style = 'Table Grid'
    table.columns[0].width = Cm(3)
    table.columns[1].width = Cm(3)
    table.columns[2].width = Cm(3)
    table.columns[3].width = Cm(3)
    table.columns[4].width = Cm(3)
    table.cell(0,0).text = '文件系统'
    table.cell(0,1).text = '总空间'
    table.cell(0,2).text = '已使用'
    table.cell(0,3).text = '可用'
    table.cell(0,4).text = '使用率'
    for fs, info in disk_info.items():
        row_cells = table.add_row().cells
        row_cells[0].text = fs
        row_cells[1].text = info['total'] + ' KB'
        row_cells[2].text = info['used'] + ' KB'
        row_cells[3].text = info['available'] + ' KB'
        row_cells[4].text = info['use%']

    # 4.3 资源使用分布图
    document.add_heading('4.3 资源使用分布图', level=3)
    mem_usage = 0
    if mem_info.get('total') and mem_info.get('used'):
        mem_usage = (int(mem_info['used']) / int(mem_info['total'])) * 100
    disk_usages = [float(info['use%'].rstrip('%')) for info in disk_info.values() if info['use%'].rstrip('%').replace('.', '').isdigit()]
    avg_disk_usage = sum(disk_usages) / len(disk_usages) if disk_usages else 0

    document.add_paragraph('以下为内存和磁盘平均使用率的可视化分布：')
    
    # Generate bar chart
    plt.figure(figsize=(6, 4))
    #labels = ['内存使用率', '磁盘平均使用率']
    labels = ['Memory usage', 'Average disk usage']
    values = [mem_usage, avg_disk_usage]
    colors = ['#36A2EB', '#FF6384']
    plt.bar(labels, values, color=colors, edgecolor=['#2E8BC0', '#D81B60'])
    plt.ylim(0, 100)
    #plt.ylabel('使用率 (%)')
    plt.ylabel('usage (%)')
    #plt.title('资源使用分布')
    plt.title('Resource usage distribution')
    plt.tight_layout()

    # Save chart to temporary file
    with tempfile.NamedTemporaryFile(suffix='.png', delete=False) as tmp_file:
        plt.savefig(tmp_file.name, format='png', dpi=300)
        tmp_file_path = tmp_file.name

    # Insert chart into document (centered)
    p = document.add_paragraph()
    p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    p.add_run().add_picture(tmp_file_path, width=Cm(15))
    plt.close()

    # Clean up temporary file
    try:
        os.remove(tmp_file_path)
    except Exception as e:
        print(f"警告: 无法删除临时文件 {tmp_file_path}: {str(e)}")

    # 4.4 关键进程
    document.add_heading('4.4 关键进程', level=3)
    processes = log_data['system_resources'].get('processes', [])
    table = document.add_table(rows=1, cols=4)
    table.style = 'Table Grid'
    table.columns[0].width = Cm(2)
    table.columns[1].width = Cm(3)
    table.columns[2].width = Cm(3)
    table.columns[3].width = Cm(7)
    table.cell(0,0).text = 'PID'
    table.cell(0,1).text = '用户'
    table.cell(0,2).text = '内存使用'
    table.cell(0,3).text = '命令'
    for proc in processes[:10]:  # Limit to top 10 processes
        row_cells = table.add_row().cells
        row_cells[0].text = proc['pid']
        row_cells[1].text = proc['user']
        row_cells[2].text = proc['vsz'] + ' KB'
        row_cells[3].text = proc['command'][:50]  # Truncate long commands

    # 5. SEL日志分析
    document.add_heading('5. 系统事件日志(SEL)分析', level=2)
    document.add_paragraph(f'日志条目总数：{log_data["sel_info"].get("Entries", "未知")}')
    document.add_paragraph(f'日志使用率：{log_data["sel_info"].get("Percent Used", "未知")}')
    
    # 5.1 事件类型统计
    document.add_heading('5.1 事件类型统计', level=3)
    table = document.add_table(rows=1, cols=2)
    table.style = 'Table Grid'
    table.columns[0].width = Cm(8)
    table.columns[1].width = Cm(7)
    table.cell(0,0).text = '事件类型'
    table.cell(0,1).text = '发生次数'
    event_types = sorted(log_data['event_summary']['by_type'].items(), key=lambda x: x[1], reverse=True)[:10]
    for event_type, count in event_types:
        row_cells = table.add_row().cells
        row_cells[0].text = translate_event_type(event_type)
        row_cells[1].text = str(count)

    # 5.2 事件类型分布图
    document.add_heading('5.2 事件类型分布图', level=3)
    document.add_paragraph('以下为事件类型分布的可视化饼图：')
    event_labels = [translate_event_type(event_type) for event_type, _ in event_types]
    event_counts = [count for _, count in event_types]

    # Generate pie chart
    plt.figure(figsize=(6, 4))
    colors = ['#FF6384', '#36A2EB', '#FFCE56', '#4BC0C0', '#9966FF',
              '#FF9F40', '#FF5733', '#C70039', '#900C3F', '#581845']
    plt.pie(event_counts, labels=event_labels, colors=colors, autopct='%1.1f%%', startangle=90)
    #plt.title('事件类型分布')
    plt.title('Event type distribution')
    plt.tight_layout()

    # Save chart to temporary file
    with tempfile.NamedTemporaryFile(suffix='.png', delete=False) as tmp_file:
        plt.savefig(tmp_file.name, format='png', dpi=300)
        tmp_file_path = tmp_file.name

    # Insert chart into document (centered)
    p = document.add_paragraph()
    p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    p.add_run().add_picture(tmp_file_path, width=Cm(15))
    plt.close()

    # Clean up temporary file
    try:
        os.remove(tmp_file_path)
    except Exception as e:
        print(f"警告: 无法删除临时文件 {tmp_file_path}: {str(e)}")

    # 5.3 电源状态变化
    if log_data['sel_events']['power_state_changes']:
        document.add_heading('5.3 电源状态变化（最近5条）', level=3)
        for event in log_data['sel_events']['power_state_changes'][-5:]:
            document.add_paragraph(f'- {event}', style='List Bullet')

    # 5.4 按日期统计事件分布图
    document.add_heading('5.4 按日期统计事件分布图', level=3)
    document.add_paragraph('以下为按日期统计的事件分布折线图：')
    def convert_to_date(date):
        try:
            return datetime.strptime(date, '%m/%d/%Y')
        except ValueError:
            return datetime(9999, 12, 31)  # 使用最大日期处理非日期字符串
    event_dates = sorted([date for date in log_data['event_summary']['by_date'].keys() if date != 'Unknown'],
                         key=convert_to_date)
    event_counts = [log_data['event_summary']['by_date'][date] for date in event_dates]

    # Generate line chart
    plt.figure(figsize=(8, 4))
    plt.plot(event_dates, event_counts, marker='o', linestyle='-', color='#36A2EB')
    #plt.xlabel('日期')
    plt.xlabel('date')
    #plt.ylabel('事件数量')
    plt.ylabel('Number of events')
    #plt.title('按日期统计事件分布')
    plt.title('Statistics of event distribution by date')
    plt.xticks(rotation=45)
    plt.grid(True)
    plt.tight_layout()

    # Save chart to temporary file
    with tempfile.NamedTemporaryFile(suffix='.png', delete=False) as tmp_file:
        plt.savefig(tmp_file.name, format='png', dpi=300)
        tmp_file_path = tmp_file.name

    # Insert chart into document (centered)
    p = document.add_paragraph()
    p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    p.add_run().add_picture(tmp_file_path, width=Cm(15))
    plt.close()

    # Clean up temporary file
    try:
        os.remove(tmp_file_path)
    except Exception as e:
        print(f"警告: 无法删除临时文件 {tmp_file_path}: {str(e)}")

    # 6. 异常事件
    document.add_heading('6. 异常事件记录', level=2)
    for error_type, errors in [
        ('电源异常', log_data['sel_events']['critical_events']),
        ('内存错误', log_data['sel_events']['memory_errors']),
        ('处理器错误', log_data['sel_events']['processor_errors'])
    ]:
        if errors:
            p = document.add_paragraph(f'{error_type}（共 {len(errors)} 条）：')
            p.runs[0].font.color.rgb = RGBColor(255, 0, 0)
            for event in errors[-5:]:  # Show last 5 events
                document.add_paragraph(f'- {event}', style='List Bullet')
        else:
            document.add_paragraph(f'未检测到{error_type}', style='List Bullet')

    # 7. 安全信息
    document.add_heading('7. 安全信息', level=2)
    table = document.add_table(rows=1, cols=3)
    table.style = 'Table Grid'
    table.columns[0].width = Cm(3)
    table.columns[1].width = Cm(5)
    table.columns[2].width = Cm(5)
    table.cell(0,0).text = '用户ID'
    table.cell(0,1).text = '用户名'
    table.cell(0,2).text = '权限'
    for user in log_data['security_info'].get('users', []):
        row_cells = table.add_row().cells
        row_cells[0].text = user['id']
        row_cells[1].text = user['name']
        row_cells[2].text = user['privilege']

    # 8. 巡检结论
    document.add_heading('8. 巡检结论', level=2)
    conclusion = []
    conclusion.append(f'设备IP: {device_ip}')
    conclusion.append(f'巡检时间: {datetime.now().strftime("%Y-%m-%d %H:%M:%S")}')
    conclusion.append(f'固件版本: {log_data["system_info"].get("Firmware Revision", "未知")}')
    
    # Power status
    if log_data.get('power_status') == 'on':
        conclusion.append('电源状态: 正常运行')
    else:
        conclusion.append('警告: 电源状态异常，请检查电源模块')

    # Memory status
    memory_issues = [slot for slot in log_data['hardware_status']['memory_slots'] if slot['status'] == 'Deasserted']
    if memory_issues:
        conclusion.append(f'警告: 检测到内存插槽异常: {[slot["slot"] for slot in memory_issues]}')
    else:
        conclusion.append(f'内存插槽: 正常，检测到 {len(log_data["hardware_status"]["memory_slots"])} 个插槽')

    # Drive status
    drive_issues = [slot for slot in log_data['hardware_status']['drive_slots'] if slot['status'] == 'Deasserted']
    if drive_issues:
        conclusion.append(f'警告: 检测到驱动器插槽异常: {[slot["slot"] for slot in drive_issues]}')
    else:
        conclusion.append(f'驱动器插槽: 正常，检测到 {len(log_data["hardware_status"]["drive_slots"])} 个插槽')

    # Critical events
    total_critical = (len(log_data['sel_events']['critical_events']) + 
                     len(log_data['sel_events']['memory_errors']) + 
                     len(log_data['sel_events']['processor_errors']))
    if total_critical:
        conclusion.append(f'严重事件: 检测到 {total_critical} 个异常事件（电源: {len(log_data["sel_events"]["critical_events"])}，'
                        f'内存: {len(log_data["sel_events"]["memory_errors"])}，处理器: {len(log_data["sel_events"]["processor_errors"])}），'
                        '需立即检查相关模块')
    else:
        conclusion.append('异常事件: 未发现严重事件')

    # Network configuration
    ip_address = log_data['network_info'].get('IP Address', '未知')
    if ip_address == device_ip:
        conclusion.append('网络配置: IP地址正确，网络连接正常')
    else:
        conclusion.append('警告: 网络配置IP与设备IP不匹配，请检查网络设置')

    # Security
    enabled_users = len([user for user in log_data['security_info'].get('users', []) if user['privilege'] != 'NO ACCESS'])
    if enabled_users > 1:
        conclusion.append(f'安全警告: 检测到 {enabled_users} 个启用用户，建议限制为仅管理员账户')
    else:
        conclusion.append('安全配置: 仅一个管理员账户启用，符合安全要求')

    # Resource usage
    if mem_info.get('total') and mem_info.get('used'):
        usage = (int(mem_info['used']) / int(mem_info['total'])) * 100
        if usage > 90:
            conclusion.append(f'警告: 内存使用率高达 {usage:.1f}%，建议优化系统资源')
        else:
            conclusion.append(f'内存使用: 正常，使用率 {usage:.1f}%')

    conclusion_text = '\n'.join(conclusion)
    p = document.add_paragraph(conclusion_text)
    p.paragraph_format.first_line_indent = Pt(20)

    document.add_page_break()

if __name__ == '__main__':
    if len(sys.argv) < 4:
        print("用法: python3 ipmi_report_generator.py <日志文件路径1> <日志文件路径2> ... <报告标题> <巡检人员>")
        sys.exit(1)

    log_paths = sys.argv[1:-2]
    report_title = sys.argv[-2]
    inspector = sys.argv[-1]

    yemei_tupian()
    top_diyiye(report_title)

    for log_path in log_paths:
        if not os.path.exists(log_path):
            print(f"错误: 日志文件 {log_path} 不存在")
            continue
        log_data = parse_ipmi_log(log_path)
        generate_report(log_data, inspector)

    report_dir = './temp/Script/ipmi_reports/'
    os.makedirs(report_dir, exist_ok=True)
    report_path = os.path.join(report_dir, f'{report_title}_IPMI巡检报告_{t1}.docx')
    try:
        document.save(report_path)
        print(f"报告已成功保存到: {report_path}")
    except Exception as e:
        print(f"保存报告失败: {str(e)}")
        import traceback
        print(traceback.format_exc())

